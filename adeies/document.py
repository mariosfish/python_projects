from docx import Document
import pandas as pd
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm , Pt

# create new document
document = Document()

#set up font
font = document.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(12)

# set up margins
sections = document.sections
for section in sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)

# Add some text 
p1 = document.add_paragraph()
p2 = document.add_paragraph()
p3 = document.add_paragraph()
p4 = document.add_paragraph()
p5 = document.add_paragraph()
p6 = document.add_paragraph()
p1 = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('ΜΗ ΥΠΗΡΕΣΙΑΚΗ').underline = True
p.add_run('Υπλγός (ΠΒ)')
p.add_run('Μιχαλόπουλος Μάριος').underline = True


# Add a heading, level 2
document.add_heading('Profile', level=2)   

# Save document
document.save(r'Document.docx')
